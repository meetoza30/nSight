import pdfplumber
import json
import requests
import re
import datetime
import time
from dateutil import parser
# from utils.generatepdf import generate_resume_pdf
from dotenv import load_dotenv
import os

load_dotenv()


def calculate_total_experience(experiences):
    total_months = 0
    now = datetime.datetime.now()
    
    for exp in experiences:
        duration_str = exp.get("duration", "")
        if not duration_str:
            continue
            
        # Clean up formats like "Aug'25" or "Aug’25" to "Aug 2025"
        duration_str = re.sub(r'[\'’](\d{2})\b', r' 20\1', duration_str)
            
        parts = re.split(r'\s*(?:-|to|–)\s*', duration_str, flags=re.IGNORECASE)
        if len(parts) == 2:
            start_str, end_str = parts[0].strip(), parts[1].strip()
            if not start_str or not end_str:
                continue
                
            try:
                start_date = parser.parse(start_str, default=datetime.datetime(now.year, 1, 1))
                if end_str.lower() in ["present", "current", "till date", "today", "now"]:
                    end_date = now
                else:
                    end_date = parser.parse(end_str, default=datetime.datetime(now.year, 1, 1))
                
                months = (end_date.year - start_date.year) * 12 + (end_date.month - start_date.month)
                if months > 0:
                    total_months += months
            except Exception:
                pass
                
    return round(total_months / 12.0, 1)

# 1. TEXT EXTRACTION — HELPERS

def _clean_cid(text):
    """Clean (cid:NNN) tokens produced by unembedded fonts and map common ligatures."""
    if not text:
        return text
    replacements = [
        ("(cid:14)", "ffi"),
        ("(cid:11)", "ff"),
        ("(cid:12)", "fi"),
        ("(cid:13)", "fl"),
    ]
    for bad, good in replacements:
        text = text.replace(bad, good)
    return re.sub(r'\(cid:\d+\)', '', text)


def _fix_mojibake(text):
    """Fix Windows-1252 characters that were double-encoded as UTF-8 (mojibake).
    Applies before text is sent to the LLM so descriptions render cleanly.
    """
    if not text:
        return text
    replacements = [
        ('\u00e2\u20ac\u201c', '\u2013'),   # â€" → en-dash
        ('\u00e2\u20ac\u201d', '\u2014'),   # â€" → em-dash
        ('\u00e2\u20ac\u2122', '\u2019'),   # â€™ → right single quote / apostrophe
        ('\u00e2\u20ac\u0153', '\u201c'),   # â€œ → left double quote
        ('\u00e2\u20ac\x9d',   '\u201d'),   # â€  → right double quote
        ('\u00e2\u20ac\xa2',   '\u2022'),   # â€¢ → bullet
        ('\u00e2\u20ac\xa6',   '\u2026'),   # â€¦ → ellipsis
        ('\u00e2\u20ac\xb0',   '\u2030'),   # â€° → per mille
        # Literal string forms that survive the decode
        ('â\u20acâ\u20ac',    '\u2013'),
        ('â\u20ac™',          "'"),
        ('â\u20ac"',          '\u2013'),
    ]
    for bad, good in replacements:
        text = text.replace(bad, good)
    return text


def _detect_column_bands(page):
    """
    Detect multi-column layout by analysing where words START horizontally (x0).

    Algorithm:
    1. Collect x0 (left edge) of every word on the page.
    2. Sort and find the largest gap between consecutive x0 values.
    3. If that gap is wide enough (>= 8% of page width AND >= 40pt), treat
       it as a column gutter and split the page there.

    Returns:
        List of (x0, x1) crop bands, one per column, left-to-right.
        None if the page is single-column.
    """
    words = page.extract_words()
    if not words or len(words) < 6:
        return None

    page_width = page.width
    if page_width < 200:
        return None

    # Collect all word left-edge positions
    x0_values = sorted(set(round(w["x0"], 1) for w in words))

    if len(x0_values) < 4:
        return None

    # Find the largest gap between consecutive x0 values
    best_gap = 0
    best_gap_left = 0
    best_gap_right = 0

    for i in range(1, len(x0_values)):
        gap = x0_values[i] - x0_values[i - 1]
        if gap > best_gap:
            best_gap = gap
            best_gap_left = x0_values[i - 1]
            best_gap_right = x0_values[i]

    # Thresholds: gap must be meaningful relative to the page
    min_gap_abs = 40          # at least 40pt wide
    min_gap_pct = 0.08        # at least 8% of page width

    if best_gap < min_gap_abs or best_gap < page_width * min_gap_pct:
        return None

    # Split point is the midpoint of the gutter
    split_x = (best_gap_left + best_gap_right) / 2.0

    # Balance check: the smaller side must contain real column content,
    # not just margin annotations (dates, links, icons).
    # Rule: minor side needs ≥ 25 words AND those words must span ≥ 40%
    # of the page height (true columns run most of the page; scattered
    # annotations don't).
    total = len(words)
    left_count = sum(1 for w in words if w["x0"] < split_x)
    right_count = total - left_count

    if left_count <= right_count:
        minor_words = [w for w in words if w["x0"] < split_x]
    else:
        minor_words = [w for w in words if w["x0"] >= split_x]

    minor_count = len(minor_words)
    if minor_count < 25:
        return None

    y_vals = [w["top"] for w in minor_words]
    y_span = max(y_vals) - min(y_vals)
    if y_span < page.height * 0.40:
        return None

    return [(0, split_x), (split_x, page_width)]


def _chars_to_lines(chars):
    """
    Group characters into lines by Y-coordinate, preserving PDF stream order.
    Returns list of (first_x0, line_text) tuples.

    - Characters within 3pt vertically are considered the same line.
    - A space is inserted when the horizontal gap exceeds 2pt.
    - first_x0 is the x0 of the very first character in the line
      (used for column assignment in multi-column pages).
    """
    if not chars:
        return []

    lines_data = []
    curr_line = []
    curr_top = None
    prev_x1 = None
    first_x0 = None

    for c in chars:
        top = c.get("top", 0)
        x0 = c.get("x0", 0)
        x1 = c.get("x1", x0)
        text = c.get("text", "")
        size = c.get("size", 10.0)

        if curr_top is None or abs(top - curr_top) <= 3.0:
            # Same line — insert space if there is a horizontal gap
            if prev_x1 is not None and (x0 - prev_x1) > 0.15 * size:
                curr_line.append(" ")
            curr_line.append(text)
            curr_top = top
            prev_x1 = max(prev_x1, x1) if prev_x1 is not None else x1
            if first_x0 is None:
                first_x0 = x0
        else:
            # New line — flush previous
            lines_data.append((first_x0, "".join(curr_line)))
            curr_line = [text]
            curr_top = top
            prev_x1 = x1
            first_x0 = x0

    if curr_line:
        lines_data.append((first_x0, "".join(curr_line)))
    return lines_data


def _extract_page_text(page):
    """
    Smart per-page extractor:
    - Auto-detects multi-column layout via word-start positions.
    - For multi-column: extracts ALL characters in stream order from the full
      page, groups them into lines, and assigns each line to left/right column
      based on where its first character starts.  This avoids crop-based word
      truncation that occurs when words straddle the column boundary.
    - For single-column: uses stream-order extraction directly.
    - Cleans (cid:NNN) tokens and ligatures.
    """
    chars = getattr(page, "chars", None)
    if not chars:
        text = page.extract_text(x_tolerance=2, y_tolerance=3)
        return _clean_cid(text) if text else ""

    bands = _detect_column_bands(page)

    if bands is None:
        # Single-column — stream-order extraction
        lines_data = _chars_to_lines(chars)
        text = "\n".join(t for _, t in lines_data)
        return _clean_cid(text)

    # Multi-column — assign each line to its column based on first_x0
    split_x = bands[0][1]       # boundary between left and right columns
    lines_data = _chars_to_lines(chars)

    left_lines = [t for fx0, t in lines_data if fx0 < split_x]
    right_lines = [t for fx0, t in lines_data if fx0 >= split_x]

    combined = "\n".join(left_lines) + "\n" + "\n".join(right_lines)
    return _clean_cid(combined.strip())


# 1. TEXT EXTRACTION

def extract_text_preserving_layout(file_path):
    """Extracts text from PDF or DOCX with auto multi-column support."""
    full_text = ""
    if file_path.lower().endswith('.docx') or file_path.lower().endswith('.doc'):
        try:
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        full_text += row_text + "\n"
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    else:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = _extract_page_text(page)
                    if page_text:
                        full_text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
    return _fix_mojibake(full_text)

# 2. LLM EXTRACTION VIA OPENROUTER

def extract_resume_json(resume_text, model="qwen/qwen3-32b"):
    """Sends the resume text to OpenRouter and returns a parsed JSON dictionary."""
    
    OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
    print(OPENROUTER_API_KEY)
    
    current_year = datetime.date.today().year
    current_month = datetime.date.today().strftime("%B")
    start = time.perf_counter()
    
    system_prompt = f"""You are a JSON-only data extraction API. Output raw valid JSON — no markdown, no explanation.

    TODAY'S DATE: {current_month} {current_year}

    CORE RULES:

    1. VERBATIM ONLY — Every value must come directly from the resume. Never invent or infer. Empty fields = "" or [].
    2. IGNORE STANDALONE PROJECTS SECTIONS — If the resume has a section headed "Projects", "Personal Projects", or "Academic Projects", DO NOT extract ANYTHING under that section heading. Do NOT attach bullet points, technologies, or dates from standalone projects to any job in the Experience section. Completely discard all content under standalone project sections.
    3. EXPERIENCE BOUNDARIES — Only extract work performed under a professional employer (jobs, internships) listed under the Experience section. Never merge or bleed content from sections outside Experience (such as standalone Projects, Volunteer, or Publications) into the experiences array. Never include personal or academic projects in experiences.
    4. COMPANY vs CLIENT vs PROJECT — "company" = employer. "client" = external org served (only if explicitly named, else ""). "project_name" = specific system/product built. EXTRACT EVERY UNIQUE EMPLOYER as a separate object in the "experiences" array. If a person had multiple roles/projects at the SAME employer, group them under that single employer's "projects" array. Do NOT merge different companies into one entry, and do NOT skip any companies.
    5. TECHNOLOGIES per project — Only extract if resume explicitly labels them (e.g. "Technologies:", "Tech Stack:"). Do NOT extract tech names from bullet point descriptions. If unlabeled, output [].
    6. PROJECT DETECTION (within an allowed Experience entry only) — Only create a separate project item when the resume explicitly names a deliverable under that job (e.g. "Project: X", a titled system/product). A plain task or achievement bullet is NOT a project. When in doubt, use ONE project item with empty project_name and put ALL of that job's bullets in description[].
    7. SKILLS — Extract ONLY from the dedicated Skills/Technical Skills section. Map categories to schema (Languages, Web Technologies, Tools, Technologies, Operating System). Keep names exactly as written.
    8. TOTAL EXPERIENCE — If resume states years explicitly, use that. Otherwise compute from earliest start date to today ("Present" = today). Round to 1 decimal. Count only professional work.
    9. DESCRIPTIONS — Copy all bullet points into "description" array. Do NOT summarize. Remove any leading bullet symbols (e.g., ◆, •, -, *, etc.).
    10. ROLE RESPONSIBILITY — Extract only if explicitly stated as "role"/"role responsibility". Do not put job titles here.
    11. JOBS WITHOUT PROJECTS — Treat entire role as single project with empty "project_name" and "client".
    12. EDUCATION GRADE — Extract the grade (CGPA, GPA, percentage, or marks) for each education entry. Look for explicit mentions like "CGPA", "GPA", "Grade", "Score", "%", or patterns like "85%", "8.5/10", "3.8 GPA". If not found, output "".
    OUTPUT SCHEMA (output ONLY this JSON, nothing else):
        {{"Name":"","Education":[{{"college":"","degree":"","graduation_year":"","grade":""}}],"Skills":{{"Programming Languages":[],"Web Technologies":[],"Tools":[],"Technologies":[],"Operating System":[]}},"Achievements":[],"Experience":{{"total_experience_years":0.0,"experiences":[{{"company":"","designation":"","duration":"","location":"","projects":[{{"client":"","project_name":"","project_span":"","technologies":[],"description":[],"role_responsibility":""}}]}}]}}}}"""



    max_retries = 3
    for attempt in range(max_retries):
        try:
            print(f"Attempt {attempt + 1}/{max_retries} to extract data...")
            response = requests.post(
                url="https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json"
                },
                data=json.dumps({
                    "model": model, 
                    "response_format": {"type": "json_object"},
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": f"Extract this resume and return ONLY a valid JSON object with no markdown:\n\n{resume_text}"}
                    ],
                    "temperature": 0.1
                })
            )
            data = response.json()
            if response.status_code != 200 or "error" in data:
                print(f"API Error ({response.status_code}): {json.dumps(data.get('error', data), indent=2)}")
                response.raise_for_status()
            print("resume extraction")
            if "usage" in data:
                print("Prompt Tokens:", data["usage"].get("prompt_tokens"))
                print("Completion Tokens:", data["usage"].get("completion_tokens"))
                print("Total Tokens:", data["usage"].get("total_tokens"))
            result_text = data['choices'][0]['message']['content'].strip()
            # print(result_text)
            end = time.perf_counter()

            print(f"LLM call took {end - start:.2f} seconds")
            
            # JSON Catcher
            start_idx = result_text.find('{')
            end_idx = result_text.rfind('}')
            
            if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
                clean_json_string = result_text[start_idx:end_idx+1]
                try:
                    parsed_json = json.loads(clean_json_string)
                    
                    # Fast validation to ensure key properties look right
                    if not isinstance(parsed_json, dict) or "Experience" not in parsed_json:
                        print(f"Format validation failed on attempt {attempt + 1}. Retrying...")
                        continue
                        
                    # --- APPLY PYTHON EXPERIENCE CALCULATION ---
                    if "Experience" in parsed_json and "experiences" in parsed_json["Experience"]:
                        calculated_years = calculate_total_experience(parsed_json["Experience"]["experiences"])
                        print("calculated_years : ", calculated_years)
                        # print()
                        # if calculated_years > 0:
                        #     parsed_json["Experience"]["total_experience_years"] = calculated_years

                    # --- CLEAN BULLET POINTS ---
                    def clean_text(text):
                        if isinstance(text, str):
                            return re.sub(r'^[◆•\-*·\s]+', '', text).strip()
                        return text

                    if "Achievements" in parsed_json and isinstance(parsed_json["Achievements"], list):
                        parsed_json["Achievements"] = [clean_text(a) for a in parsed_json["Achievements"] if clean_text(a)]
                        
                    if "Experience" in parsed_json and isinstance(parsed_json["Experience"], dict) and "experiences" in parsed_json["Experience"]:
                        for exp in parsed_json["Experience"]["experiences"]:
                            if "projects" in exp and isinstance(exp["projects"], list):
                                for proj in exp["projects"]:
                                    if "description" in proj and isinstance(proj["description"], list):
                                        proj["description"] = [clean_text(d) for d in proj["description"] if clean_text(d)]

                    return parsed_json
                except json.JSONDecodeError as e:
                    print(f"JSON Parsing Error on attempt {attempt + 1}: {e}")
                    if attempt == max_retries - 1:
                        print(f"Raw output that failed:\n{clean_json_string}")
                        return None
            else:
                print(f"Error: Could not find {{}} brackets on attempt {attempt + 1}.")
                if attempt == max_retries - 1:
                    print(f"Raw output:\n{result_text}")
                    return None
                
        except Exception as e:
            if hasattr(e, 'response') and e.response is not None:
                print(f"Extraction failed on attempt {attempt + 1}: {e}\nResponse: {e.response.text}")
            else:
                print(f"Extraction failed on attempt {attempt + 1}: {e}")
            if attempt == max_retries - 1:
                return None
            time.sleep(2)
            
    return None

# 3. MAIN RUNNER

def process_resume(file_path):
    print(f"Reading file: {file_path}...")
    
    full_text = extract_text_preserving_layout(file_path)
    if not full_text:
        print("Empty or unreadable file")
        return
    print("text", full_text)
    print("Extracting data via LLM...")
    # extracted_data = "meet"
    extracted_data = extract_resume_json(full_text)
    
    if extracted_data:
        print("\n--- EXTRACTION RESULT ---")
        print(json.dumps(extracted_data, indent=4, ensure_ascii=False))
        
        # Generate the PDF (Uncomment when ready)
        # candidate_name = extracted_data.get("Name", "Unknown")
        # pdf_name = candidate_name.replace(" ", "_") + "_NcircleCV.pdf"
        # generate_resume_pdf(extracted_data, pdf_name)
        # print(f"Generated PDF: {pdf_name}")

if __name__ == "__main__":
    resume_paths = [
        # "./Resumes/2pg2.pdf",
        # "./Resumes/1quarter.pdf",
        # "./Resumes/03_Data_Scientist_Rahul_Mehta.pdf",
        # "./Resumes/Aman_Babu_s_Resume__Copy_ (1).pdf",
        # "./Resumes/civil.pdf",
        # "./Resumes/icons_resume.pdf",
        # "./Resumes/1manyinfo2.pdf",
        # "./Resumes/1_2cols_pic.pdf",
        # "./Resumes/1_page_many_info.pdf",
        # "./Resumes/3_page.pdf",
    ]

    for pdf_path in resume_paths:
        process_resume(pdf_path)